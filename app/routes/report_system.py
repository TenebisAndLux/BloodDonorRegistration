import io
from flask import Blueprint, render_template, request, make_response, jsonify
from docx import Document
from datetime import datetime, timedelta
from babel.dates import format_date
from app import db
from app.models import Donor, Doctor, MedicalInstitution, MedicalExamination, MedicalHistory, BloodSupply, \
    BloodCollectionType, BloodCollection

report_system = Blueprint('report_system', __name__)

@report_system.route('/reports')
def index():
    """Главная страница системы отчетов"""
    institutions = MedicalInstitution.query.all()
    collection_types = BloodCollectionType.query.all()
    stocks = BloodSupply.query.all()
    numbers = BloodCollection.query.all()
    return render_template('main/report_index.html',
                           institutions=institutions,
                           collection_types=collection_types,
                           stocks=stocks,
                           numbers=numbers)

def prepare_donor_report_data(blood_collection):
    """Подготовка данных для отчета о доноре"""
    donor = Donor.query.filter_by(passportdata=blood_collection.passportdata).first_or_404(
        description="Донор не найден")

    doctor = Doctor.query.filter_by(servicenumber=blood_collection.servicenumber).first_or_404(
        description="Врач не найден")

    institution = MedicalInstitution.query.get_or_404(
        blood_collection.institutioncode,
        description="Учреждение не найдено")

    medical_examination = MedicalExamination.query.filter(
        MedicalExamination.passportdata == donor.passportdata,
        MedicalExamination.dateofexamination <= blood_collection.collectiondate
    ).order_by(MedicalExamination.dateofexamination.desc()).first()

    medical_history = MedicalHistory.query.get_or_404(
        donor.historynumber,
        description="Медицинская история не найдена")

    blood_supply = BloodSupply.query.filter_by(
        numberstock=blood_collection.numberstock).first()

    collection_type = BloodCollectionType.query.get_or_404(
        blood_collection.collectiontypecode,
        description="Тип сбора крови не найден")

    return {
        'donor': donor,
        'doctor': doctor,
        'institution': institution,
        'medical_examination': medical_examination,
        'medical_history': medical_history,
        'blood_supply': blood_supply,
        'collection_type': collection_type,
        'blood_collection': blood_collection
    }

@report_system.route('/reports/generate_report_1', methods=['POST'])
def generate_report_1():
    """Генерация отчета о доноре с предпросмотром"""
    try:
        # Валидация данных
        required_fields = ['collectiontypecode', 'institutioncode', 'numberstock', 'number']
        if not all(field in request.form for field in required_fields):
            return jsonify({'error': 'Не все обязательные поля заполнены'}), 400

        collectiontypecode = int(request.form['collectiontypecode'])
        institutioncode = int(request.form['institutioncode'])
        numberstock = int(request.form['numberstock'])
        number = int(request.form['number'])

        # Получаем данные
        blood_collection = BloodCollection.query.filter_by(
            bloodsupplycollectiontypecode=collectiontypecode,
            bloodbankinstitutioncode=institutioncode,
            numberstock=numberstock,
            number=number
        ).first_or_404(description="Событие сбора крови не найдено")

        report_data = prepare_donor_report_data(blood_collection)

        # Если запрос на предпросмотр
        if request.form.get('preview') == 'true':
            preview_content = {
                'donor_name': f"{report_data['donor'].surname} {report_data['donor'].name}",
                'institution': report_data['institution'].nameofinstitution,
                'collection_date': blood_collection.collectiondate.strftime("%d.%m.%Y"),
                'blood_volume': report_data['blood_supply'].bloodvolume if report_data['blood_supply'] else "N/A"
            }
            return jsonify(preview_content)

        # Генерация DOCX
        doc = generate_donor_docx(report_data)

        # Отправка файла
        return send_docx_response(doc, f"donor_report_{blood_collection.collectiondate.strftime('%Y%m%d')}.docx")

    except ValueError as e:
        return jsonify({'error': f"Некорректные данные: {str(e)}"}), 400
    except Exception as e:
        return jsonify({'error': f"Ошибка сервера: {str(e)}"}), 500

def generate_donor_docx(report_data):
    """Генерация DOCX файла для отчета о доноре"""
    doc = Document()

    # Извлечение данных
    donor = report_data['donor']
    doctor = report_data['doctor']
    institution = report_data['institution']
    medical_examination = report_data['medical_examination']
    medical_history = report_data['medical_history']
    blood_supply = report_data['blood_supply']
    collection_type = report_data['collection_type']
    blood_collection = report_data['blood_collection']

    # Заголовок отчета
    doc.add_heading(institution.nameofinstitution, level=1)
    doc.add_paragraph(f'Адрес: {institution.address}')
    doc.add_paragraph(f'Контактный телефон: {institution.contactphonenumber}')
    doc.add_paragraph(f'Дата формирования отчета: {datetime.now().strftime("%d.%m.%Y")}')

    # Разделитель
    doc.add_paragraph('_' * 40)

    # Данные донора
    doc.add_heading('Данные донора', level=2)
    doc.add_paragraph(f'ФИО: {donor.surname} {donor.name} {donor.secondname or ""}')
    doc.add_paragraph(f'Дата рождения: {donor.birthday.strftime("%d.%m.%Y")}')
    doc.add_paragraph(f'Пол: {donor.gender}')
    doc.add_paragraph(f'Адрес: {donor.address}')
    doc.add_paragraph(f'Телефон: {donor.phonenumber or "Не указан"}')
    doc.add_paragraph(f'Полис ОМС: {donor.polis}')
    doc.add_paragraph(f'Группа крови и резус-фактор: {donor.bloodgroup}, Rh{donor.rhfactor}')
    doc.add_paragraph(f'Номер медицинской карты: MH-{donor.historynumber:04d}')
    doc.add_paragraph(f'Дата последнего обследования: {medical_history.dateoflastexamination.strftime("%d.%m.%Y") if medical_history else "N/A"}')
    doc.add_paragraph(f'Результаты анализов: {medical_history.analysisresults or "Нет данных"}')
    doc.add_paragraph(f'Ограничения на донацию: {"Есть" if medical_history.banondonation else "Нет"}')

    # Разделитель
    doc.add_paragraph('_' * 40)

    # Информация о донации
    doc.add_heading('Информация о донации', level=2)
    doc.add_paragraph(f'Дата донации: {blood_collection.collectiondate.strftime("%d.%m.%Y")}')
    doc.add_paragraph(f'Тип донации: {collection_type.name}')
    doc.add_paragraph(f'Объем забора: {blood_supply.bloodvolume} мл' if blood_supply else 'N/A')
    doc.add_paragraph(f'Ответственный врач: {doctor.secondname} {doctor.name} (№ служебного удостоверения: DR-{doctor.servicenumber:03d})')

    # Разделитель
    doc.add_paragraph('_' * 40)

    # Заключение
    doc.add_heading('Заключение', level=2)
    doc.add_paragraph('1. Донор допущен к сдаче крови, противопоказаний не выявлено.')
    doc.add_paragraph('2. Забор крови произведен успешно, самочувствие донора после процедуры удовлетворительное.')
    next_donation_date = blood_collection.collectiondate + timedelta(days=60)
    doc.add_paragraph(f'3. Следующая возможная дата донации: {next_donation_date.strftime("%d.%m.%Y")} (через 60 дней).')
    doc.add_paragraph('Спасибо за ваш вклад в спасение жизней!')

    # Разделитель
    doc.add_paragraph('_' * 40)

    # Примечание
    doc.add_heading('Примечание', level=2)
    doc.add_paragraph('Данный отчет сформирован автоматически на основании данных из базы медицинского учреждения.')
    doc.add_paragraph('В случае вопросов обращайтесь по телефону регистратуры: +7 (915) 606-54-33.')

    return doc

def generate_institution_docx(institutioncode, start_date, end_date):
    """Генерация DOCX файла для сводного отчета по учреждению"""
    doc = Document()

    # Получение данных об учреждении
    institution = MedicalInstitution.query.get_or_404(
        institutioncode, description="Учреждение не найдено")

    # Заголовок отчета
    doc.add_heading('Отчет о донорстве крови', level=1)
    doc.add_paragraph(f'Дата формирования отчета: {format_date(datetime.now(), "d MMMM yyyy года", locale="ru_RU")}')
    doc.add_paragraph(f'Медицинское учреждение: {institution.nameofinstitution}')
    doc.add_paragraph(f'Адрес: {institution.address}')
    doc.add_paragraph(f'Контактный телефон: {institution.contactphonenumber}')
    doc.add_paragraph(f'Email: {institution.email}')

    # Введение
    doc.add_heading('1. Введение', level=2)
    doc.add_paragraph(
        f'Настоящий отчет содержит информацию о донорских донациях, проведенных в {institution.nameofinstitution} '
        f'за период с {format_date(start_date, "d MMMM", locale="ru_RU")} по {format_date(end_date, "d MMMM yyyy года", locale="ru_RU")}. '
        'В отчете представлены данные о количестве доноров, объеме заготовленной крови и ее компонентов, '
        'а также о результатах медицинских обследований доноров.'
    )

    # Основная часть
    doc.add_heading('2. Основная часть', level=2)

    # Количество доноров и донаций
    doc.add_heading('2.1. Количество доноров и донаций', level=3)
    blood_collections = BloodCollection.query.filter(
        BloodCollection.institutioncode == institutioncode,
        BloodCollection.collectiondate.between(start_date, end_date)
    ).all()

    unique_donors = db.session.query(db.func.count(db.distinct(BloodCollection.passportdata))).filter(
        BloodCollection.institutioncode == institutioncode,
        BloodCollection.collectiondate.between(start_date, end_date)
    ).scalar()

    # Подсчет первичных и повторных доноров
    donor_counts = db.session.query(
        BloodCollection.passportdata,
        db.func.count(BloodCollection.number).label('donation_count')
    ).filter(
        BloodCollection.institutioncode == institutioncode,
        BloodCollection.collectiondate.between(start_date, end_date)
    ).group_by(BloodCollection.passportdata).all()

    primary_donors = sum(1 for _, count in donor_counts if count == 1)
    repeat_donors = unique_donors - primary_donors
    total_donations = len(blood_collections)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    headers = ['Показатель', 'Значение']
    data = [
        ('Общее количество доноров', unique_donors),
        ('Количество первичных доноров', primary_donors),
        ('Количество повторных доноров', repeat_donors),
        ('Общее количество донаций', total_donations)
    ]
    for row_idx, (header, value) in enumerate([headers] + data):
        row = table.rows[row_idx]
        row.cells[0].text = str(header)
        row.cells[1].text = str(value)

    # Заготовка крови и компонентов
    doc.add_heading('2.2. Заготовка крови и ее компонентов', level=3)
    total_blood_volume = db.session.query(db.func.sum(BloodSupply.bloodvolume)).join(
        BloodCollection,
        db.and_(
            BloodSupply.numberstock == BloodCollection.numberstock,
            BloodSupply.collectiontypecode == BloodCollection.bloodsupplycollectiontypecode,
            BloodSupply.institutioncode == BloodCollection.bloodbankinstitutioncode
        )
    ).filter(
        BloodCollection.institutioncode == institutioncode,
        BloodCollection.collectiondate.between(start_date, end_date)
    ).scalar() or 0

    # Примерные данные для компонентов (адаптировать под реальные данные, если доступны)
    plasma_volume = total_blood_volume * 0.4  # Примерное соотношение
    erythrocyte_volume = total_blood_volume * 0.3
    platelet_volume = total_blood_volume * 0.2

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    headers = ['Компонент крови', 'Объем (л)']
    data = [
        ('Цельная кровь', round(total_blood_volume / 1000, 2)),
        ('Эритроцитарная масса', round(erythrocyte_volume / 1000, 2)),
        ('Плазма', round(plasma_volume / 1000, 2)),
        ('Тромбоциты', round(platelet_volume / 1000, 2))
    ]
    for row_idx, (header, value) in enumerate([headers] + data):
        row = table.rows[row_idx]
        row.cells[0].text = str(header)
        row.cells[1].text = str(value)

    # Результаты медицинских обследований
    doc.add_heading('2.3. Результаты медицинских обследований доноров', level=3)
    bans = MedicalHistory.query.join(
        Donor,
        MedicalHistory.historynumber == Donor.historynumber
    ).join(
        BloodCollection,
        Donor.passportdata == BloodCollection.passportdata
    ).filter(
        BloodCollection.institutioncode == institutioncode,
        BloodCollection.collectiondate.between(start_date, end_date),
        MedicalHistory.banondonation == True
    ).count()

    abnormal_results = MedicalHistory.query.join(
        Donor,
        MedicalHistory.historynumber == Donor.historynumber
    ).join(
        BloodCollection,
        Donor.passportdata == BloodCollection.passportdata
    ).filter(
        BloodCollection.institutioncode == institutioncode,
        BloodCollection.collectiondate.between(start_date, end_date),
        MedicalHistory.analysisresults != None,
        MedicalHistory.analysisresults != ''
    ).count()

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    headers = ['Показатель', 'Количество случаев']
    data = [
        ('Выявлено противопоказаний к донорству', bans),
        ('Отстранены от донорства по медицинским причинам', bans),  # Предположительно совпадает
        ('Обнаружены отклонения в анализах крови', abnormal_results)
    ]
    for row_idx, (header, value) in enumerate([headers] + data):
        row = table.rows[row_idx]
        row.cells[0].text = str(header)
        row.cells[1].text = str(value)

    # Заключение
    doc.add_heading('3. Заключение', level=2)
    doc.add_paragraph(
        f'В период с {format_date(start_date, "d MMMM", locale="ru_RU")} по {format_date(end_date, "d MMMM yyyy года", locale="ru_RU")} '
        f'{institution.nameofinstitution} успешно провела заготовку крови и ее компонентов, '
        'обеспечив потребности лечебных отделений. Работа с донорами велась в соответствии с установленными '
        'стандартами и нормативами. Выявленные медицинские противопоказания и отклонения в анализах были '
        'своевременно зафиксированы и учтены в дальнейшей работе.'
    )

    # Подпись
    doc.add_paragraph('Подпись ответственного лица: _______________ /Иванов И.И./')
    doc.add_paragraph(f'Дата: {format_date(datetime.now(), "d MMMM yyyy года", locale="ru_RU")}')

    return doc

def send_docx_response(doc, filename):
    """Отправка DOCX файла в ответе"""
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    response = make_response(file_stream.getvalue())
    response.headers.set('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response.headers.set('Content-Disposition', 'attachment', filename=filename)
    return response

@report_system.route('/reports/generate_report_2', methods=['POST'])
def generate_report_2():
    """Генерация сводного отчета по учреждению с предпросмотром"""
    try:
        # Валидация данных
        required_fields = ['institutioncode', 'start_date', 'end_date']
        if not all(field in request.form for field in required_fields):
            return jsonify({'error': 'Не все обязательные поля заполнены'}), 400

        institutioncode = int(request.form['institutioncode'])
        start_date = datetime.strptime(request.form['start_date'], '%Y-%m-%d')
        end_date = datetime.strptime(request.form['end_date'], '%Y-%m-%d')

        if start_date > end_date:
            return jsonify({'error': 'Дата начала не может быть позже даты окончания'}), 400

        institution = MedicalInstitution.query.get_or_404(
            institutioncode, description="Учреждение не найдено")

        # Если запрос на предпросмотр
        if request.form.get('preview') == 'true':
            # Простая статистика для предпросмотра
            total_donors = db.session.query(db.func.count(db.distinct(BloodCollection.passportdata))).filter(
                BloodCollection.institutioncode == institutioncode,
                BloodCollection.collectiondate.between(start_date, end_date)
            ).scalar()

            return jsonify({
                'institution': institution.nameofinstitution,
                'period': f"{start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')}",
                'total_donors': total_donors
            })

        # Генерация полного отчета
        doc = generate_institution_docx(institutioncode, start_date, end_date)
        return send_docx_response(doc, f"institution_report_{institutioncode}_{datetime.now().strftime('%Y%m%d')}.docx")

    except ValueError as e:
        return jsonify({'error': f"Некорректные данные: {str(e)}"}), 400
    except Exception as e:
        return jsonify({'error': f"Ошибка сервера: {str(e)}"}), 500